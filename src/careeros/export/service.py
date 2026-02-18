# =========================
# FILE: src/careeros/export/service.py
# =========================
from __future__ import annotations

import glob
import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from careeros.generation.schema import ApplicationPackage
from careeros.guardrails.schema import ValidationReport
from careeros.parsing.schema import EvidenceProfile


def _ensure_dir(p: str | Path) -> Path:
    path = Path(p)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _latest(pattern: str) -> Optional[str]:
    files = sorted(glob.glob(pattern))
    return files[-1] if files else None


def _safe_job_hint(job_path: str) -> str:
    # turn outputs/jobs/job_post_v1_xxx.json into a short hint
    base = Path(job_path).stem
    base = re.sub(r"[^a-zA-Z0-9_]+", "_", base)
    return base[-40:] if len(base) > 40 else base


def load_json(path: str) -> dict:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def load_package(path: str) -> ApplicationPackage:
    return ApplicationPackage.model_validate(load_json(path))


def load_report(path: str) -> ValidationReport:
    return ValidationReport.model_validate(load_json(path))


def load_profile(path: str) -> EvidenceProfile:
    return EvidenceProfile.model_validate(load_json(path))


def render_docx(
    pkg: ApplicationPackage,
    profile: Optional[EvidenceProfile],
    out_path: Path,
) -> Path:
    """
    Create a recruiter-friendly DOCX from ApplicationPackage.
    Phase 0: minimal formatting, deterministic content.
    """
    doc = Document()

    doc.add_heading("CareerOS Application Package", level=1)

    # Metadata
    doc.add_paragraph(f"Run ID: {getattr(pkg, 'run_id', '')}")
    doc.add_paragraph(f"Profile: {getattr(pkg, 'profile_path', '')}")
    doc.add_paragraph(f"Job: {getattr(pkg, 'job_path', '')}")

    doc.add_paragraph("")  # spacer

    # Full tailored resume
    tr = getattr(pkg, "tailored_resume", None)
    if tr:
        doc.add_heading("Tailored Resume (Full Draft)", level=2)
        if getattr(tr, "headline", None):
            doc.add_paragraph(tr.headline)
        if getattr(tr, "professional_summary", None):
            doc.add_paragraph(tr.professional_summary)
        skills = getattr(tr, "core_skills", []) or []
        if skills:
            doc.add_paragraph("Core Skills: " + ", ".join(skills))
        for x in (getattr(tr, "experience_highlights", []) or []):
            doc.add_paragraph(x, style="List Bullet")
        for x in (getattr(tr, "projects_highlights", []) or []):
            doc.add_paragraph(x, style="List Bullet")

    # Tailored bullets
    doc.add_heading("Tailored Resume Bullets", level=2)
    bullets = getattr(pkg, "bullets", []) or []
    if bullets:
        for b in bullets:
            txt = getattr(b, "text", "")
            if txt:
                doc.add_paragraph(txt, style="List Bullet")
    else:
        doc.add_paragraph("(No bullets found)")

    # Evidence skills (optional)
    ev = getattr(pkg, "evidence_skills", None)
    if ev:
        doc.add_heading("Evidence Skills Used", level=2)
        doc.add_paragraph(", ".join(ev))

    # Cover letter
    cover = getattr(pkg, "cover_letter", None)
    if cover:
        doc.add_heading("Cover Letter Draft", level=2)
        subject = getattr(cover, "subject", None)
        if subject:
            doc.add_paragraph(f"Subject: {subject}")
        paras = getattr(cover, "paragraphs", []) or []
        for p in paras:
            if p:
                doc.add_paragraph(p)

    # QA stubs
    qa = getattr(pkg, "qa_stubs", None)
    if isinstance(qa, dict) and qa:
        doc.add_heading("Application Q/A Stubs", level=2)
        for q, a in qa.items():
            if q:
                doc.add_paragraph(f"Q: {q}")
            if a:
                doc.add_paragraph(f"A: {a}")
            doc.add_paragraph("")

    _ensure_dir(out_path.parent)
    doc.save(str(out_path))
    return out_path


def render_pdf_from_package(
    pkg: ApplicationPackage,
    out_path: Path,
) -> Path:
    """
    Generate a PDF directly from ApplicationPackage using ReportLab.
    No external DOCX->PDF converters required.
    """
    _ensure_dir(out_path.parent)
    c = canvas.Canvas(str(out_path), pagesize=LETTER)
    width, height = LETTER

    x = 50
    y = height - 50
    line_h = 14

    def write_line(text: str, bold: bool = False):
        nonlocal y
        if y < 70:
            c.showPage()
            y = height - 50
        c.setFont("Helvetica-Bold" if bold else "Helvetica", 11 if bold else 10)
        c.drawString(x, y, text[:110])  # simple truncation for Phase 0
        y -= line_h

    write_line("CareerOS Application Package", bold=True)
    write_line(f"Run ID: {getattr(pkg, 'run_id', '')}")
    write_line(f"Profile: {getattr(pkg, 'profile_path', '')}")
    write_line(f"Job: {getattr(pkg, 'job_path', '')}")
    y -= line_h

    tr = getattr(pkg, "tailored_resume", None)
    if tr:
        write_line("Tailored Resume (Full Draft)", bold=True)
        if getattr(tr, "headline", None):
            write_line(tr.headline)
        if getattr(tr, "professional_summary", None):
            write_line(getattr(tr, "professional_summary"))
        skills = getattr(tr, "core_skills", []) or []
        if skills:
            write_line("Core Skills: " + ", ".join(skills[:12]))
        for x in (getattr(tr, "experience_highlights", []) or []):
            write_line(f"• {x}")
        for x in (getattr(tr, "projects_highlights", []) or []):
            write_line(f"• {x}")
        y -= line_h

    write_line("Tailored Resume Bullets", bold=True)
    bullets = getattr(pkg, "bullets", []) or []
    if bullets:
        for b in bullets:
            txt = getattr(b, "text", "")
            if txt:
                write_line(f"• {txt}")
    else:
        write_line("(No bullets found)")

    y -= line_h
    cover = getattr(pkg, "cover_letter", None)
    if cover:
        write_line("Cover Letter Draft", bold=True)
        subject = getattr(cover, "subject", None)
        if subject:
            write_line(f"Subject: {subject}")
        paras = getattr(cover, "paragraphs", []) or []
        for p in paras:
            if not p:
                continue
            # naive wrap
            words = p.split()
            buf = []
            for w in words:
                buf.append(w)
                if len(" ".join(buf)) > 95:
                    write_line(" ".join(buf))
                    buf = []
            if buf:
                write_line(" ".join(buf))
            y -= 2

    qa = getattr(pkg, "qa_stubs", None)
    if isinstance(qa, dict) and qa:
        y -= line_h
        write_line("Application Q/A Stubs", bold=True)
        for q, a in qa.items():
            if q:
                write_line(f"Q: {q}")
            if a:
                write_line(f"A: {a}")
            y -= 2

    c.save()
    return out_path


def export_latest_validated_package(
    *,
    package_path: Optional[str] = None,
    validation_report_path: Optional[str] = None,
    profile_path: Optional[str] = None,
    out_dir: str = "exports/submissions",
    export_docx: bool = True,
    export_pdf: bool = True,
) -> dict:
    """
    Phase 0: pick latest artifacts if paths not provided.
    Enforce: only export if validation_report.status == 'pass'.
    Returns export metadata dict (API will wrap into ExportResult).
    """
    if package_path is None:
        package_path = _latest("exports/packages/application_package_v1_*.json")
        if not package_path:
            raise FileNotFoundError("No application_package_v1_*.json found under exports/packages/")

    if validation_report_path is None:
        validation_report_path = _latest("outputs/guardrails/validation_report_v1_*.json")
        if not validation_report_path:
            raise FileNotFoundError("No validation_report_v1_*.json found under outputs/guardrails/")

    pkg = load_package(package_path)
    report = load_report(validation_report_path)

    if report.status != "pass":
        raise ValueError(f"Export blocked: validation_report status is '{report.status}'. Fix package and re-validate.")

    # profile is optional; include if you have it
    if profile_path is None:
        profile_path = _latest("outputs/profile/profile_v1_*.json")
    profile = load_profile(profile_path) if profile_path else None

    run_id = getattr(pkg, "run_id", "") or "run_unknown"
    job_path = getattr(pkg, "job_path", "") or ""
    job_hint = _safe_job_hint(job_path) if job_path else "job"

    application_id = f"app_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"

    base_dir = Path(out_dir) / f"{application_id}_{job_hint}"
    _ensure_dir(base_dir)

    docx_path = None
    pdf_path = None

    if export_docx:
        docx_path = str(render_docx(pkg, profile, base_dir / "application_package.docx"))

    if export_pdf:
        pdf_path = str(render_pdf_from_package(pkg, base_dir / "application_package.pdf"))

    return {
        "application_id": application_id,
        "run_id": run_id,
        "package_path": package_path,
        "validation_report_path": validation_report_path,
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "status": "exported",
    }
